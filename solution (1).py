from docx import Document
from docx.shared import Pt
from docx.text import font

option = int(input("Вариант:"))

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(14)
p = document.add_paragraph()
p.add_run('Введение').italic = True
document.add_paragraph()
document.add_paragraph().add_run('Цель курсовой работы: '
                                 'рассчитать показатели экономической эффективности изготовления ').italic = True
document.add_paragraph().add_run('Задачи курсовой работы:').italic = True
for i in (['определить исходные данные для расчета себестоимости изделия;',
           'рассчитать полную себестоимость изделия;', 'построить диаграмму структуры полной себестоимости изделия;',
           'определить оптовую цену изделия;', 'построить график безубыточности производства;',
           'построить график безубыточности производства;',
           'определить годовой объем выпуска изделий из условия безубыточности производства;',
           'определить годовой фонд заработной платы основных рабочих;',
           'определить численность основных производственных рабочих;',
           'определить срок окупаемости дополнительных капиталовложений.']):
    p = document.add_paragraph()
    p.add_run(i).italic = True
    p.style = 'List Number'

document.add_page_break()
document.add_paragraph().add_run('1 Исходные данные').italic = True
document.add_paragraph().add_run('Данные для расчета стоимости материалов представлены в таблице 1.1.').italic = True
list_table_1 = [['№\nп/п', 'Наименование материала', 'Единица измерения', 'Количество',
                 'Цена за единицу измерения, руб.'],
                ['1', 'Припой ПОС-61', 'кг', str(round(0.001 * option + 0.05, 3)), '2400'],
                ['2', 'Флюс ЛТИ-120', 'л', str(round(0.001 * option + 0.02, 3)), '780'],
                ['3', 'Спирт', 'л', '0.004', '110'],
                ['4', 'Лак УР-231', 'л', '0.03', '689']]

table = document.add_table(rows=5, cols=5)
table.style = 'Table Grid'
table.italic = True

for i in range(5):
    hdr_cells = table.rows[i].cells
    for j in range(5):
        hdr_cells[j].text = list_table_1[i][j]
        hdr_cells[j].italic = True

document.save('test.docx')
